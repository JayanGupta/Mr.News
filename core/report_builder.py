"""
report_builder.py — Builds DOCX and PDF reports from Markdown analysis sections.
Handles tables, headings, charts, and branding for Mr.News by Jayan Gupta.
New structure: Cover → Executive Summary → Expert Research → Plain English Guide → Sources
"""

import os
import io
import re
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image as RLImage, PageBreak,
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


# ─── Colour Palette ───────────────────────────────────────────────────────────
BRAND_GREEN = (16, 185, 129)        # #10b981
BRAND_BLUE  = (59, 130, 246)        # #3b82f6
BRAND_DARK  = (15, 23, 42)          # #0f172a
BRAND_ACCENT = (245, 158, 11)       # #f59e0b  (amber)
BRAND_PURPLE = (124, 58, 237)       # #7c3aed

RL_GREEN  = colors.Color(16/255, 185/255, 129/255)
RL_BLUE   = colors.Color(59/255, 130/255, 246/255)
RL_DARK   = colors.Color(15/255, 23/255, 42/255)
RL_ACCENT = colors.Color(245/255, 158/255, 11/255)
RL_LIGHT  = colors.Color(248/255, 250/255, 252/255)
RL_PURPLE = colors.Color(124/255, 58/255, 237/255)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _parse_md_tables(text: str):
    """Extract markdown tables from text as list of (header_row, data_rows)."""
    pattern = re.compile(
        r"(\|[^\n]+\|)\n\|[-| :]+\|\n((?:\|[^\n]+\|\n?)+)",
        re.MULTILINE,
    )
    tables = []
    for m in pattern.finditer(text):
        header = [c.strip() for c in m.group(1).split("|") if c.strip()]
        rows_raw = m.group(2).strip().splitlines()
        rows = []
        for row in rows_raw:
            cells = [c.strip() for c in row.split("|") if c.strip()]
            if cells:
                rows.append(cells)
        tables.append((header, rows, m.start(), m.end()))
    return tables


def _clean_md(text: str) -> str:
    """Strip markdown syntax for plain-text contexts."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"#{1,6}\s*", "", text)
    return text.strip()


# ─── DOCX Builder ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour in DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_docx_heading(doc: Document, text: str, level: int):
    text = _clean_md(text)
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(*BRAND_GREEN)
        run.font.size = Pt(22)
    elif level == 2:
        run.font.color.rgb = RGBColor(*BRAND_BLUE)
        run.font.size = Pt(16)
    else:
        run.font.color.rgb = RGBColor(*BRAND_ACCENT)
        run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _add_docx_table(doc: Document, header: list, rows: list):
    all_rows = [header] + rows
    cols = max(len(r) for r in all_rows)
    # Pad short rows
    all_rows = [r + [""] * (cols - len(r)) for r in all_rows]

    tbl = doc.add_table(rows=len(all_rows), cols=cols)
    tbl.style = "Table Grid"

    for ci, cell_text in enumerate(all_rows[0]):
        cell = tbl.cell(0, ci)
        cell.text = cell_text
        _set_cell_bg(cell, "10b981")
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    for ri, row in enumerate(all_rows[1:], start=1):
        bg = "f8fafc" if ri % 2 == 0 else "e2e8f0"
        for ci, cell_text in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = cell_text
            _set_cell_bg(cell, bg)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()


def _md_section_to_docx(doc: Document, text: str):
    """Convert a Markdown section to DOCX elements."""
    tables = _parse_md_tables(text)
    table_spans = {(t[2], t[3]): (t[0], t[1]) for t in tables}

    pos = 0
    lines = list(enumerate(text.splitlines()))
    line_positions = []
    cur = 0
    for i, line in enumerate(text.splitlines()):
        line_positions.append(cur)
        cur += len(line) + 1  # +1 for newline

    skip_until = -1

    for i, line in enumerate(text.splitlines()):
        line_start = line_positions[i]
        line_end = line_start + len(line)

        # Skip lines that are inside a table we already rendered
        if i <= skip_until:
            continue

        # Check if this line starts a table
        in_table = False
        for (ts, te), (hdr, rows) in table_spans.items():
            if line_start >= ts and line_end <= te:
                in_table = True
                # Render table once (when first line of table encountered)
                if line_start == ts or abs(line_start - ts) < 5:
                    _add_docx_table(doc, hdr, rows)
                    # figure out how many lines to skip
                    table_text = text[ts:te]
                    skip_until = i + len(table_text.splitlines()) - 1
                break

        if in_table:
            continue

        stripped = line.strip()
        if not stripped or stripped.startswith("|---"):
            continue

        if stripped.startswith("### "):
            _add_docx_heading(doc, stripped[4:], 3)
        elif stripped.startswith("## "):
            _add_docx_heading(doc, stripped[3:], 2)
        elif stripped.startswith("# "):
            _add_docx_heading(doc, stripped[2:], 1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:])
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped.startswith("```"):
            pass  # skip code fences
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)
            p.paragraph_format.space_after = Pt(4)


def _add_inline_runs(para, text: str):
    """Handle **bold** and *italic* inline markdown."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(*BRAND_GREEN)
        else:
            para.add_run(part)


def _add_section_divider(doc: Document, title: str, emoji: str, color: tuple):
    """Add a styled section divider page."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run(f"{emoji}  {title}")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    doc.add_paragraph()


def build_docx(topic: str, sections: dict, sources: list) -> bytes:
    """Build and return DOCX bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover Page ────────────────────────────────────────────────────────────
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(100)
    run = cover_title.add_run("📰 Mr.News")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*BRAND_GREEN)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Expert-Level Deep Research Report  •  Built by Jayan")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(*BRAND_BLUE)
    sub_run.italic = True

    doc.add_paragraph()
    topic_p = doc.add_paragraph()
    topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = topic_p.add_run(f"Research Topic: {topic}")
    tr.font.size = Pt(20)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(*BRAND_ACCENT)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y  %H:%M')}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(100, 116, 139)

    # Report type label
    label_p = doc.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_before = Pt(20)
    lr = label_p.add_run("Comprehensive Research Report with Expert Analysis & Plain English Guide")
    lr.font.size = Pt(9)
    lr.font.color.rgb = RGBColor(148, 163, 184)
    lr.italic = True

    doc.add_page_break()

    # ── Table of Contents placeholder ────────────────────────────────────────
    _add_docx_heading(doc, "📑 Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Expert Research Report",
        "   2.1 What Is It?",
        "   2.2 Core Features & Capabilities",
        "   2.3 How It Works — Technical Deep Dive",
        "   2.4 Who Uses It & Real-World Applications",
        "   2.5 How to Get Started",
        "   2.6 Competitive Landscape",
        "   2.7 Strengths & Limitations",
        "   2.8 Latest News & Developments",
        "   2.9 Expert Verdict & Recommendations",
        "3. Plain English Guide",
        "4. Sources & References",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 116, 139)
        p.paragraph_format.space_after = Pt(2)

    # ── Executive Summary ─────────────────────────────────────────────────────
    _add_section_divider(doc, "Executive Summary", "📋", BRAND_GREEN)
    _md_section_to_docx(doc, sections.get("summary", ""))

    # ── Expert Research ──────────────────────────────────────────────────────
    _add_section_divider(doc, "Expert Research Report", "🔬", BRAND_BLUE)
    _md_section_to_docx(doc, sections.get("expert_research", ""))

    # ── Plain English Guide ───────────────────────────────────────────────────
    _add_section_divider(doc, "Plain English Guide", "📖", BRAND_PURPLE)
    _md_section_to_docx(doc, sections.get("plain_english", ""))

    # ── Sources ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_docx_heading(doc, "🌐 Sources & References", 1)
    p_note = doc.add_paragraph()
    note_run = p_note.add_run(
        f"This report was compiled from {len(sources)} web sources scraped and analyzed by Mr.News AI."
    )
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(100, 116, 139)
    note_run.italic = True
    doc.add_paragraph()

    for i, src in enumerate(sources[:25], 1):
        p = doc.add_paragraph(style="List Number")
        run_title = p.add_run(f"{src['title']}: ")
        run_title.bold = True
        run_url = p.add_run(src["url"])
        run_url.font.color.rgb = RGBColor(*BRAND_BLUE)
        run_url.font.size = Pt(9)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("— End of Report —")
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(100, 116, 139)
    fr.italic = True

    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_p.add_run(
        "Generated by Mr.News AI Research Reporter  •  Built by Jayan  •  "
        "Powered by Gemini 2.0 + DuckDuckGo"
    )
    disc_run.font.size = Pt(8)
    disc_run.font.color.rgb = RGBColor(148, 163, 184)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── PDF Builder ──────────────────────────────────────────────────────────────

def _get_pdf_styles():
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "MrTitle", parent=styles["Title"],
        fontSize=36, textColor=RL_GREEN,
        spaceAfter=6, alignment=TA_CENTER, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "MrSubtitle", parent=styles["Normal"],
        fontSize=13, textColor=RL_BLUE,
        spaceAfter=4, alignment=TA_CENTER, fontName="Helvetica-Oblique",
    ))
    styles.add(ParagraphStyle(
        "MrTopic", parent=styles["Normal"],
        fontSize=20, textColor=RL_ACCENT,
        spaceAfter=6, alignment=TA_CENTER, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "MrH1", parent=styles["Heading1"],
        fontSize=18, textColor=RL_GREEN,
        spaceBefore=14, spaceAfter=5, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "MrH2", parent=styles["Heading2"],
        fontSize=14, textColor=RL_BLUE,
        spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "MrH3", parent=styles["Heading3"],
        fontSize=11, textColor=RL_ACCENT,
        spaceBefore=8, spaceAfter=3, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "MrBody", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=4,
        fontName="Helvetica", alignment=TA_JUSTIFY,
    ))
    styles.add(ParagraphStyle(
        "MrBullet", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=3,
        leftIndent=20, fontName="Helvetica",
        bulletIndent=10,
    ))
    styles.add(ParagraphStyle(
        "MrSectionDivider", parent=styles["Title"],
        fontSize=24, textColor=RL_GREEN,
        spaceBefore=80, spaceAfter=10, alignment=TA_CENTER, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "MrFooter", parent=styles["Normal"],
        fontSize=8, textColor=colors.Color(0.58, 0.64, 0.72),
        spaceAfter=2, alignment=TA_CENTER, fontName="Helvetica-Oblique",
    ))
    return styles


def _md_to_pdf_elements(text: str, styles):
    elements = []
    tables = _parse_md_tables(text)
    table_spans = [(t[2], t[3], t[0], t[1]) for t in tables]

    line_positions = []
    cur = 0
    for line in text.splitlines():
        line_positions.append(cur)
        cur += len(line) + 1

    rendered_tables = set()
    skip_until = -1

    for i, line in enumerate(text.splitlines()):
        if i <= skip_until:
            continue

        line_start = line_positions[i]

        in_table = False
        for ts, te, hdr, rows in table_spans:
            if ts <= line_start < te:
                in_table = True
                if ts not in rendered_tables:
                    rendered_tables.add(ts)
                    all_rows = [hdr] + rows
                    cols = max(len(r) for r in all_rows)
                    all_rows = [r + [""] * (cols - len(r)) for r in all_rows]

                    col_width = (A4[0] - 3*cm) / cols
                    tdata = []
                    for ri, row in enumerate(all_rows):
                        tdata.append([Paragraph(_clean_md(c), styles["MrBody"]) for c in row])

                    pdf_table = Table(tdata, colWidths=[col_width]*cols, repeatRows=1)
                    ts_style = TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), RL_GREEN),
                        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE",   (0, 0), (-1, 0), 9),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [RL_LIGHT, colors.white]),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.Color(0.8, 0.8, 0.8)),
                        ("FONTSIZE",  (0, 1), (-1, -1), 8),
                        ("VALIGN",    (0, 0), (-1, -1), "TOP"),
                        ("TOPPADDING",    (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
                    ])
                    pdf_table.setStyle(ts_style)
                    elements.append(pdf_table)
                    elements.append(Spacer(1, 8))
                    table_text = text[ts:te]
                    skip_until = i + len(table_text.splitlines()) - 1
                break

        if in_table:
            continue

        stripped = line.strip()
        if not stripped or stripped.startswith("|---") or stripped.startswith("```"):
            continue

        def esc(t):
            return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if stripped.startswith("### "):
            elements.append(Paragraph(esc(_clean_md(stripped[4:])), styles["MrH3"]))
        elif stripped.startswith("## "):
            elements.append(Paragraph(esc(_clean_md(stripped[3:])), styles["MrH2"]))
        elif stripped.startswith("# "):
            elements.append(Paragraph(esc(_clean_md(stripped[2:])), styles["MrH1"]))
        elif stripped.startswith(("- ", "* ")):
            elements.append(Paragraph(f"• {esc(_clean_md(stripped[2:]))}", styles["MrBullet"]))
        elif re.match(r"^\d+\.", stripped):
            content = re.sub(r"^\d+\.\s*", "", stripped)
            elements.append(Paragraph(f"{esc(_clean_md(content))}", styles["MrBullet"]))
        else:
            elements.append(Paragraph(esc(_clean_md(stripped)), styles["MrBody"]))

    return elements


def build_pdf(topic: str, sections: dict, sources: list) -> bytes:
    """Build and return PDF bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )
    styles = _get_pdf_styles()
    story = []

    # ── Cover Page ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 1.5*inch))
    story.append(Paragraph("📰 Mr.News", styles["MrTitle"]))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph("Expert-Level Deep Research Report  •  Built by Jayan", styles["MrSubtitle"]))
    story.append(Spacer(1, 0.4*inch))
    story.append(HRFlowable(width="100%", thickness=2, color=RL_GREEN))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(f"Research Topic: {topic}", styles["MrTopic"]))
    story.append(Spacer(1, 0.2*inch))
    date_str = datetime.datetime.now().strftime("%B %d, %Y  %H:%M")
    story.append(Paragraph(f"Generated: {date_str}", styles["MrBody"]))
    story.append(Spacer(1, 0.15*inch))
    story.append(Paragraph(
        "Comprehensive Research Report with Expert Analysis &amp; Plain English Guide",
        styles["MrFooter"],
    ))
    story.append(PageBreak())

    # ── Executive Summary ─────────────────────────────────────────────────────
    story.append(Paragraph("📋 Executive Summary", styles["MrSectionDivider"]))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_GREEN))
    story.append(Spacer(1, 12))
    story.extend(_md_to_pdf_elements(sections.get("summary", ""), styles))
    story.append(PageBreak())

    # ── Expert Research ──────────────────────────────────────────────────────
    story.append(Paragraph("🔬 Expert Research Report", styles["MrSectionDivider"]))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_BLUE))
    story.append(Spacer(1, 12))
    story.extend(_md_to_pdf_elements(sections.get("expert_research", ""), styles))
    story.append(PageBreak())

    # ── Plain English Guide ───────────────────────────────────────────────────
    divider_style = ParagraphStyle(
        "MrPurpleDivider", parent=styles["MrSectionDivider"],
        textColor=RL_PURPLE,
    )
    story.append(Paragraph("📖 Plain English Guide", divider_style))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_PURPLE))
    story.append(Spacer(1, 12))
    story.extend(_md_to_pdf_elements(sections.get("plain_english", ""), styles))
    story.append(PageBreak())

    # ── Sources ───────────────────────────────────────────────────────────────
    story.append(Paragraph("🌐 Sources &amp; References", styles["MrH1"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        f"<i>This report was compiled from {len(sources)} web sources scraped and analyzed by Mr.News AI.</i>",
        styles["MrFooter"],
    ))
    story.append(Spacer(1, 8))
    for i, src in enumerate(sources[:25], 1):
        story.append(Paragraph(
            f"{i}. <b>{src['title']}</b><br/>"
            f"<font color='#3b82f6' size='8'>{src['url']}</font>",
            styles["MrBody"],
        ))

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5*inch))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_GREEN))
    story.append(Spacer(1, 8))
    story.append(Paragraph("— End of Report —", styles["MrFooter"]))
    story.append(Paragraph(
        "Generated by Mr.News AI Research Reporter  •  Built by Jayan  •  "
        "Powered by Gemini 2.0 + DuckDuckGo",
        styles["MrFooter"],
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()
